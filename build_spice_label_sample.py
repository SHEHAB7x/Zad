from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("generated")
FULL_PATH = OUT_DIR / "qasr_altawabel_product_labels_full.docx"

PAGE_W = 8.27
PAGE_H = 11.69
MARGIN = 0.55
CONTENT_W = PAGE_W - (MARGIN * 2)
CARD_H = 4.85

FONT_UI = "Arial"
FONT_TITLE = "Arial"
FONT_PRODUCT = "Arial"

COLORS = {
    "ink_green": "153E2F",
    "spice_red": "8F2F24",
    "gold": "C99A2E",
    "warm_paper": "FFFDF6",
}

BRAND_TITLE = "\u0642\u0635\u0631 \u0627\u0644\u062a\u0648\u0627\u0628\u0644"
DECORATIVE_RULE = "\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640  \u2726  \u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640\u0640"
DECORATIVE_MARK = "\u2726"

PRODUCTS = [
    "\u062d\u0628\u0647\u0627\u0646",
    "\u0642\u0631\u0646\u0641\u0644",
    "\u062d\u0628 \u0627\u0644\u0631\u0634\u0627\u062f",
    "\u062d\u0631\u0645\u0644",
    "\u0645\u0644\u062d \u0644\u064a\u0645\u0648\u0646",
    "\u062f\u0648\u0645 \u0645\u0637\u062d\u0648\u0646",
    "\u0641\u0644\u0641\u0644 \u0627\u0628\u064a\u0636",
    "\u0645\u0633\u062a\u0643\u0629",
    "\u0644\u0628\u0627\u0646 \u0627\u0644\u0630\u0643\u0631",
    "\u0642\u0633\u0637 \u0647\u0646\u062f\u064a",
    "\u0646\u062e\u0648\u0647 \u0647\u0646\u062f\u064a\u0647",
    "\u062f\u0648\u0646\u0643\u0627\u0631",
    "\u0628\u0644\u062d \u0627\u0644\u0637\u064a\u0628",
    "\u062c\u0648\u0632 \u0647\u0646\u062f",
    "\u0641\u0644\u0641\u0644 \u0645\u0637\u062d\u0648\u0646",
    "\u0643\u0645\u0648\u0646 \u0645\u0637\u062d\u0648\u0646",
    "\u0643\u0633\u0628\u0631\u0647 \u0645\u0637\u062d\u0648\u0646\u0629",
    "\u062d\u0644\u0628\u0629 \u0645\u0637\u062d\u0648\u0646\u0629",
    "\u063a\u0644\u0629 \u0641\u062c\u0644",
    "\u063a\u0644\u0629 \u062c\u0631\u062c\u064a\u0631",
    "\u063a\u0644\u0629 \u0628\u0627\u0645\u064a\u0629",
    "\u063a\u0644\u0629 \u0645\u0644\u0648\u062e\u064a\u0629",
    "\u0628\u0642\u062f\u0648\u0646\u0633",
    "\u0643\u0631\u0641\u0633",
    "\u062c\u0646\u0632\u0628\u064a\u0644",
    "\u0642\u0631\u0641\u0647 \u0645\u0637\u062d\u0648\u0646\u0647",
    "\u0642\u0631\u0641\u0647 \u0633\u064a\u062c\u0627\u0631",
    "\u0628\u0627\u0628\u0648\u0646\u062c",
    "\u0643\u0645\u0648\u0646 \u0627\u0633\u0648\u062f",
    "\u062d\u0628\u0629 \u0627\u0644\u0628\u0631\u0643\u0629",
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        edge_data = kwargs[edge]
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key in ("sz", "val", "color", "space"):
            if key in edge_data:
                element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_cell_margins(cell, top=120, start=260, bottom=120, end=260):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_fixed_width(table, width_inches):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))

    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)

    grid = tbl.tblGrid
    for grid_col in grid.gridCol_lst:
        grid.remove(grid_col)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(int(width_inches * 1440)))
    grid.append(grid_col)

    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(width_inches)
            tc_w = cell._tc.get_or_add_tcPr().tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width_inches * 1440)))


def set_paragraph_rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, font_name):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:cs"), font_name)
    lang = run._element.rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run._element.rPr.append(lang)
    lang.set(qn("w:bidi"), "ar-SA")


def add_text(paragraph, text, font, size, color, bold=False):
    set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_card(cell, product):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, COLORS["warm_paper"])
    set_cell_margins(cell)
    set_cell_borders(
        cell,
        top={"val": "single", "sz": 18, "color": COLORS["gold"], "space": 0},
        bottom={"val": "single", "sz": 18, "color": COLORS["gold"], "space": 0},
        left={"val": "single", "sz": 18, "color": COLORS["gold"], "space": 0},
        right={"val": "single", "sz": 18, "color": COLORS["gold"], "space": 0},
    )

    for paragraph in cell.paragraphs:
        paragraph.text = ""

    title = cell.paragraphs[0]
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    add_text(title, BRAND_TITLE, FONT_TITLE, 28, COLORS["ink_green"], bold=True)

    sep = cell.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(34)
    add_text(sep, DECORATIVE_RULE, FONT_UI, 13, COLORS["gold"], bold=False)

    product_p = cell.add_paragraph()
    product_p.paragraph_format.space_before = Pt(0)
    product_p.paragraph_format.space_after = Pt(30)
    add_text(product_p, product, FONT_PRODUCT, 54, COLORS["spice_red"], bold=True)

    low_sep = cell.add_paragraph()
    low_sep.paragraph_format.space_before = Pt(0)
    low_sep.paragraph_format.space_after = Pt(0)
    add_text(low_sep, DECORATIVE_MARK, FONT_UI, 15, COLORS["gold"], bold=False)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(PAGE_W)
    section.page_height = Inches(PAGE_H)
    section.top_margin = Inches(MARGIN)
    section.bottom_margin = Inches(MARGIN)
    section.left_margin = Inches(MARGIN)
    section.right_margin = Inches(MARGIN)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_UI
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT_UI)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def zero_trailing_paragraph_after(table):
    paragraph = table._element.getnext()
    if paragraph is None or paragraph.tag != qn("w:p"):
        return
    p_pr = paragraph.get_or_add_pPr()
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "1")
    spacing.set(qn("w:lineRule"), "exact")


def add_product_page(doc, top_product, bottom_product):
    table = doc.add_table(rows=2, cols=1)
    set_table_fixed_width(table, CONTENT_W)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    zero_trailing_paragraph_after(table)

    for row in table.rows:
        row.height = Inches(CARD_H)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    add_card(table.cell(0, 0), top_product)
    add_card(table.cell(1, 0), bottom_product)


def build_full_document():
    if len(PRODUCTS) % 2 != 0:
        raise ValueError("Product count must be even for two cards per page.")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    for index in range(0, len(PRODUCTS), 2):
        if index:
            doc.add_page_break()
        add_product_page(doc, PRODUCTS[index], PRODUCTS[index + 1])

    doc.core_properties.title = f"{BRAND_TITLE} - \u0645\u0644\u0635\u0642\u0627\u062a \u0627\u0644\u0645\u0646\u062a\u062c\u0627\u062a"
    doc.core_properties.subject = "Arabic printable spice shop product labels"
    doc.save(FULL_PATH)
    print(FULL_PATH.resolve())


if __name__ == "__main__":
    build_full_document()
